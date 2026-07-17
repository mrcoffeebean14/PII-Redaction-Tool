import re
import docx

class DocxRedactor:
    def __init__(self, detector, fake_gen):
        self.detector = detector
        self.fake_gen = fake_gen
        self.global_mappings = {}

    def _scan_text_for_pii(self, text, is_dob_col=False):
        # Triggers PII detection and registers matches in FakeGenerator
        if not text.strip():
            return
        detections = self.detector.detect(text, is_table_dob_column=is_dob_col)
        for d in detections:
            entity_text = d['text'].strip()
            entity_type = d['type']
            if entity_text and len(entity_text) > 1:
                self.fake_gen.get_fake(entity_text, entity_type)

    def _check_table_dob_columns(self, table):
        # Identifies columns with DOB-related header words
        dob_cols = set()
        if len(table.rows) == 0:
            return dob_cols
        header_row = table.rows[0]
        dob_keywords = ["dob", "birth", "born", "age", "date of birth"]
        for col_idx, cell in enumerate(header_row.cells):
            cell_text = cell.text.lower()
            if any(k in cell_text for k in dob_keywords):
                dob_cols.add(col_idx)
        return dob_cols

    def scan_document(self, doc_path):
        # Pass 1: Scans all elements (paragraphs, tables, headers, footers) to build name mappings
        doc = docx.Document(doc_path)
        
        for p in doc.paragraphs:
            self._scan_text_for_pii(p.text)
            
        for table in doc.tables:
            dob_cols = self._check_table_dob_columns(table)
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    is_dob_col = (row_idx > 0) and (col_idx in dob_cols)
                    for p in cell.paragraphs:
                        self._scan_text_for_pii(p.text, is_dob_col)
                        
        for section in doc.sections:
            for p in section.header.paragraphs:
                self._scan_text_for_pii(p.text)
            for table in section.header.tables:
                dob_cols = self._check_table_dob_columns(table)
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        is_dob_col = (row_idx > 0) and (col_idx in dob_cols)
                        for p in cell.paragraphs:
                            self._scan_text_for_pii(p.text, is_dob_col)
                            
            for p in section.footer.paragraphs:
                self._scan_text_for_pii(p.text)
            for table in section.footer.tables:
                dob_cols = self._check_table_dob_columns(table)
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        is_dob_col = (row_idx > 0) and (col_idx in dob_cols)
                        for p in cell.paragraphs:
                            self._scan_text_for_pii(p.text, is_dob_col)
                            
        self.global_mappings = self.fake_gen.mapping

    def _replace_text_in_runs(self, paragraph, matches):
        # Replaces text boundaries within runs from right to left to keep formats intact
        for start, end, replacement in matches:
            runs_boundaries = []
            current_idx = 0
            for run in paragraph.runs:
                run_len = len(run.text)
                runs_boundaries.append((current_idx, current_idx + run_len, run))
                current_idx += run_len
                
            first_run_idx = -1
            last_run_idx = -1
            for idx, (run_start, run_end, run) in enumerate(runs_boundaries):
                if run_start <= start < run_end:
                    first_run_idx = idx
                if run_start < end <= run_end:
                    last_run_idx = idx
                    
            if first_run_idx == -1 or last_run_idx == -1:
                continue
                
            if first_run_idx == last_run_idx:
                run_start, run_end, run = runs_boundaries[first_run_idx]
                offset_start = start - run_start
                offset_end = end - run_start
                run.text = run.text[:offset_start] + replacement + run.text[offset_end:]
            else:
                f_start, f_end, f_run = runs_boundaries[first_run_idx]
                offset_start = start - f_start
                f_run.text = f_run.text[:offset_start] + replacement
                
                for idx in range(first_run_idx + 1, last_run_idx):
                    _, _, mid_run = runs_boundaries[idx]
                    mid_run.text = ""
                    
                l_start, l_end, l_run = runs_boundaries[last_run_idx]
                offset_end = end - l_start
                l_run.text = l_run.text[offset_end:]

    def redact_paragraph(self, paragraph):
        # Finds PII strings in paragraph text and triggers run-level replacement
        text = paragraph.text
        if not text or not self.global_mappings:
            return
            
        sorted_keys = sorted(self.global_mappings.keys(), key=len, reverse=True)
        sorted_keys = [k for k in sorted_keys if len(k) > 1]
        
        if not sorted_keys:
            return
            
        escaped_keys = [re.escape(k) for k in sorted_keys]
        pattern = re.compile("|".join(escaped_keys))
        
        matches = []
        for m in pattern.finditer(text):
            start = m.start()
            end = m.end()
            orig = m.group()
            replacement = self.global_mappings[orig]
            matches.append((start, end, replacement))
            
        if not matches:
            return
            
        matches_sorted = sorted(matches, key=lambda x: x[0], reverse=True)
        self._replace_text_in_runs(paragraph, matches_sorted)

    def redact_document(self, input_path, output_path):
        # Pass 2: Applies redactions to paragraphs, cells, headers/footers, and saves document
        doc = docx.Document(input_path)
        
        for p in doc.paragraphs:
            self.redact_paragraph(p)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self.redact_paragraph(p)
                        
        for section in doc.sections:
            for p in section.header.paragraphs:
                self.redact_paragraph(p)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self.redact_paragraph(p)
                            
            for p in section.footer.paragraphs:
                self.redact_paragraph(p)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self.redact_paragraph(p)
                            
        doc.save(output_path)
